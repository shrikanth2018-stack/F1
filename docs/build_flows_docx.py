#!/usr/bin/env python3
"""
1stOne F1 — branded .docx generator for "Document 1: App Flows".

Renders docs/01-app-flows.docx (dark theme) with:
  - a dark cover page carrying the real "1stOne.in" wordmark + title
  - a provenance note as the first line of the document body
  - a clickable, auto-updating Table of Contents (Word TOC field, outline levels)
  - a callout-colour legend
  - flows grouped into three audience parts (page break before each)
  - two inline box-chain diagrams (cycle clock; order journey)
  - brand-coloured headings (sky #38bdf8) on Tahoma, near-white body text
  - styled bullet lists with light-green (logo) labels
  - shaded "Why" / "Standout" callout cards
  - footer with page numbers

Content is factual and code-derived (mirrors 01-app-flows.md).
Run:  python3 docs/build_flows_docx.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
LOGO = os.path.join(HERE, "_logo_real.png")
if not os.path.exists(LOGO):
    LOGO = os.path.join(ROOT, "assets", "icon.png")
OUT = os.path.join(HERE, "01-app-flows.docx")

# ── Brand palette (from src/theme/index.ts) ──────────────────────
SKY      = RGBColor(0x38, 0xBD, 0xF8)
MINT     = RGBColor(0x4E, 0xCD, 0xC4)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
FONT     = "Tahoma"
BODY       = RGBColor(0xEC, 0xEC, 0xEC)   # near-white body text on dark
LIGHTGREEN = RGBColor(0xC5, 0xE0, 0xA0)   # logo "One" green — accents / bullet labels
DARKTEXT   = RGBColor(0x22, 0x22, 0x22)   # text inside light callout cards
FOOTGREY   = RGBColor(0xAA, 0xAA, 0xAA)   # footer / fine print on dark
CONTENT_W  = 6.6                          # printable width (8.5" - 2×0.9" margins, minus a hair)

# ── low-level helpers ────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_page_bg(doc, hex_color):
    bg = OxmlElement("w:background"); bg.set(qn("w:color"), hex_color)
    doc.element.insert(0, bg)
    doc.settings.element.append(OxmlElement("w:displayBackgroundShape"))

def set_update_fields(doc):
    """Make Word rebuild field results (the TOC) when the document opens."""
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true")
    doc.settings.element.append(el)

def set_outline_level(paragraph, level):
    pPr = paragraph._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), str(level))
    pPr.append(ol)

def add_page_number_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("1stOne — App Flows   ·   ")
    run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = FOOTGREY
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r = p.add_run(); r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = FOOTGREY
    r._r.append(b); r._r.append(instr); r._r.append(e)

def style_base(doc):
    n = doc.styles["Normal"]
    n.font.name = FONT; n.font.size = Pt(10.5); n.font.color.rgb = BODY
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.12

# ── building blocks ──────────────────────────────────────────────
def part(doc, text):
    """Audience-group divider — own page, mint heading, thick rule, outline lvl 0."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(19); run.font.color.rgb = MINT
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "2"); bottom.set(qn("w:color"), "4ECDC4")
    pbdr.append(bottom); pPr.append(pbdr)
    set_outline_level(p, 0)
    return p

def heading(doc, text, size=15, color=SKY, space_before=14, outline=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(size); run.font.color.rgb = color
    if outline is not None:
        set_outline_level(p, outline)
    return p

def rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "38BDF8")
    pbdr.append(bottom); pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(2)

def body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.name = FONT; p.runs[0].font.color.rgb = BODY
    return p

def bullets(doc, items):
    for label, rest in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
        if label:
            r = p.add_run(label); r.font.name = FONT; r.bold = True; r.font.color.rgb = LIGHTGREEN
        r2 = p.add_run(rest); r2.font.name = FONT; r2.font.color.rgb = BODY

def callout(doc, kind, text):
    fill = "EAFBF9" if kind == "why" else "E8F6FE"
    label = "Why it works this way" if kind == "why" else "What makes it stand out"
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.cell(0, 0); set_cell_bg(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders"); left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18"); left.set(qn("w:space"), "0")
    left.set(qn("w:color"), "4ECDC4" if kind == "why" else "38BDF8")
    borders.append(left); tcPr.append(borders)
    cell.paragraphs[0].text = ""
    lr = cell.paragraphs[0].add_run(label.upper())
    lr.font.name = FONT; lr.font.size = Pt(8); lr.bold = True
    lr.font.color.rgb = RGBColor(0x0b, 0x7a, 0x73) if kind == "why" else RGBColor(0x0a, 0x6b, 0x9b)
    bp = cell.add_paragraph(text)
    bp.runs[0].font.name = FONT; bp.runs[0].font.size = Pt(10); bp.runs[0].font.color.rgb = DARKTEXT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def diagram_chain(doc, stages, fill_hex):
    """Horizontal box-chain: shaded stage chips separated by arrows."""
    n = len(stages); cols = 2 * n - 1
    arrow_w = 0.30
    stage_w = (CONTENT_W - arrow_w * (n - 1)) / n
    tbl = doc.add_table(rows=1, cols=cols)
    tbl.autofit = False; tbl.allow_autofit = False
    row = tbl.rows[0]; row.height = Pt(22); row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    for i in range(cols):
        w = Inches(stage_w if i % 2 == 0 else arrow_w)
        tbl.columns[i].width = w
        cell = tbl.cell(0, i); cell.width = w
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(2); para.paragraph_format.space_after = Pt(2)
        if i % 2 == 0:
            set_cell_bg(cell, fill_hex)
            r = para.add_run(stages[i // 2])
            r.font.name = FONT; r.font.size = Pt(8.5); r.bold = True; r.font.color.rgb = WHITE
        else:
            r = para.add_run("→")
            r.font.name = FONT; r.font.size = Pt(12); r.font.color.rgb = BODY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def flow(doc, title, flow_text, why, standout, sub_bullets=None, diagram=None):
    heading(doc, title, size=15, color=SKY)
    rule(doc)
    body(doc, flow_text)
    if sub_bullets:
        bullets(doc, sub_bullets)
    if diagram:
        cap, stages, fillhex = diagram
        cp = doc.add_paragraph(); cr = cp.add_run(cap)
        cr.font.name = FONT; cr.font.size = Pt(9); cr.italic = True; cr.font.color.rgb = FOOTGREY
        cp.paragraph_format.space_after = Pt(3)
        diagram_chain(doc, stages, fillhex)
    callout(doc, "why", why)
    callout(doc, "standout", standout)

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    run._r.append(b); run._r.append(instr); run._r.append(sep)
    ph = p.add_run("Right-click here and choose “Update Field” to build the contents.")
    ph.font.name = FONT; ph.font.size = Pt(9.5); ph.italic = True; ph.font.color.rgb = FOOTGREY
    endrun = p.add_run()
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    endrun._r.append(e)

# ════════════════════════════════════════════════════════════════
doc = Document()
style_base(doc)

# ---- COVER (dark) ----
sec = doc.sections[0]
sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
set_page_bg(doc, "151515")
set_update_fields(doc)

for _ in range(4):
    doc.add_paragraph()
if os.path.exists(LOGO):
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lp.add_run().add_picture(LOGO, width=Inches(4.2))
doc.add_paragraph()
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2r = t2.add_run("App Flows"); t2r.font.name = FONT; t2r.font.size = Pt(32); t2r.font.color.rgb = WHITE
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("A plain-English working guide to what the app does, how, and why")
sr.font.name = FONT; sr.font.size = Pt(12); sr.font.color.rgb = MINT
for _ in range(6):
    doc.add_paragraph()
note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note.add_run("Document 1 of a series")
nr.font.name = FONT; nr.font.size = Pt(9); nr.font.color.rgb = FOOTGREY

# ---- BODY ----
doc.add_section(WD_SECTION.NEW_PAGE)
body_sec = doc.sections[1]
body_sec.top_margin = Inches(0.8); body_sec.bottom_margin = Inches(0.8)
body_sec.left_margin = Inches(0.9); body_sec.right_margin = Inches(0.9)
add_page_number_footer(body_sec)

# Row 1 — provenance note (kept at the very top, as requested).
prov = doc.add_paragraph()
pr = prov.add_run("Prepared from a direct reading of the application's source code. "
                  "Where the code and older comments disagreed, the code was treated as the source of truth.")
pr.font.name = FONT; pr.font.size = Pt(9.5); pr.italic = True; pr.font.color.rgb = LIGHTGREEN
prov.paragraph_format.space_after = Pt(12)

# Contents (not outlined, so it doesn't list itself)
ch = doc.add_paragraph(); chr_ = ch.add_run("Contents")
chr_.font.name = FONT; chr_.font.size = Pt(18); chr_.font.color.rgb = MINT
ch.paragraph_format.space_after = Pt(4)
add_toc(doc)

# Intro page
doc.add_page_break()
heading(doc, "What this is", size=16, color=SKY, space_before=2, outline=1)
body(doc,
     "1stOne is one product that serves four kinds of people from a single login: customers "
     "(who order meals and groceries), staff (kitchen and packing), drivers / hub operators "
     "(last-mile delivery), and admins (the back-office). It runs on Android and iOS phones, the "
     "same app also runs in a web browser, and there is a public marketing website at 1stone.in. "
     "Which screens you see is decided automatically the moment you sign in. The flows below are "
     "grouped by who uses them — Customer first, then Staff & Operations, then Admin & Platform. "
     "Everything here is drawn from the actual code; technology notes appear inline, only where "
     "they help explain a function.")
# Callout legend
leg = doc.add_paragraph(); lgr = leg.add_run("How to read each section — every flow ends with two boxes:")
lgr.font.name = FONT; lgr.font.size = Pt(10); lgr.italic = True; lgr.font.color.rgb = BODY
leg.paragraph_format.space_before = Pt(8); leg.paragraph_format.space_after = Pt(4)
callout(doc, "why", "The mint box — the business or user benefit behind the design.")
callout(doc, "standout", "The blue box — how this design uniquely serves that intent.")

# ════════════ PART 1 — CUSTOMER ════════════
part(doc, "Part 1 — For the Customer")

flow(doc, "1.  Registration & Login",
     "A customer signs in with their phone number and a one-time SMS code (OTP) — there are no "
     "passwords (identity is handled by Supabase Auth, the app's backend service). A returning user "
     "lands straight on their home screen; a brand-new user does one onboarding step that captures "
     "name, first delivery address, and map location together in a single save, so an account is "
     "never half-created. The app knows who you are — customer, staff, driver, or admin — from a few "
     "signed values inside your login token, so it routes you instantly with no extra lookups. A "
     "referral code from a shared link is applied automatically once you're signed in. This works "
     "the same on the phone app and in the web browser.",
     "Phone + OTP removes the biggest sign-up drop-off (passwords) and ties every account to one "
     "verifiable identity. Deciding the role from the login token means the right screens appear "
     "with zero delay and no chance of a customer seeing staff tools.",
     "One app and one login serve all four roles — there is no separate staff or admin app. A driver "
     "who is also staff even reuses the customer app for deliveries while keeping staff permissions, "
     "all driven by that single signed token.")

flow(doc, "2.  Delivery Cycles & What You Can Order",
     "Everything you can order belongs to a delivery cycle — a fixed daily window the store runs, "
     "set up by the admin and typically aligned to meal times. Each cycle has its own order cutoff, "
     "a kitchen-start time, and a delivery-start time. You don't get random, on-demand runs; your "
     "order joins the next cycle whose cutoff hasn't passed, and the home screen shows a "
     "“Today / Tomorrow” badge so you know which run it will catch. Both meals (Food) and "
     "groceries (Essentials) are tagged to a cycle, so when a food item and an essential share the "
     "same cycle they arrive together in one delivery — and you pay the delivery fee only once. Miss "
     "a cutoff and the item simply moves to the next eligible run.",
     "Batching orders into a few daily cycles is what makes fresh, home-style food viable — the "
     "kitchen cooks to a known list at a known time instead of reacting to random orders all day, "
     "which controls quality, waste, and cost.",
     "Because food and groceries can ride the same cycle, one trip serves both needs — and the "
     "cutoff / kitchen / delivery clock is enforced by the server for every order, so the whole "
     "operation moves as predictable batches rather than chaotic one-off runs.",
     diagram=("How one cycle moves through the day:",
              ["Order cutoff", "Kitchen starts", "Delivery"], "4ECDC4"))

flow(doc, "3.  Orders — End to End",
     "On the home screen the customer browses Food and Essentials, grouped by delivery cycle, with "
     "“Today / Tomorrow” badges the server calculates from each cycle's cutoff — never the "
     "phone's clock. Items go into a cart saved on the device. At checkout the app sends only the "
     "cart and chosen address to the server, which returns the binding price, delivery date(s), "
     "tax, and delivery fee. When the customer taps Pay, the app echoes that exact quote back; the "
     "server re-checks it down to the paise, and if anything has drifted (a price changed, a cutoff "
     "passed) it refuses and asks the customer to review the fresh total — no charge happens on a "
     "stale price. Payment is by wallet (instant) or Razorpay (cards / UPI / net-banking); a "
     "Razorpay order is confirmed after its payment signature is verified, with the gateway's "
     "webhook as a backup so it confirms even if the app closes. (In the web app, online card / UPI "
     "payment isn't available — web orders pay by wallet; the phone app handles online payments.) "
     "A single order may split into one delivery per cycle but stays one “order” to the "
     "customer, who can cancel the whole order (with an automatic wallet refund) within the allowed "
     "window and before the kitchen locks the first cycle.",
     "The customer always sees exactly what they will pay and when it will arrive, and that promise "
     "can't quietly change between preview and payment. Two payment paths plus a webhook backstop "
     "mean an order is never “paid but lost.”",
     "The server is the single source of truth for money and timing — the same logic prices the "
     "preview and the charge, and a paise-level “drift tripwire” blocks any mismatch before "
     "a rupee moves. A one-time order key makes duplicate orders impossible.",
     diagram=("Your order's journey:",
              ["Confirmed", "Ready", "Packed", "Dispatched", "Delivered"], "38BDF8"))

flow(doc, "4.  Wallet & Top-Up",
     "Each customer has a prepaid wallet shown on their profile. Topping up creates a Razorpay "
     "payment; the wallet is credited only after the payment is verified on the server — never from "
     "the app directly. Minimum and maximum top-up amounts are admin-set. The wallet pays for orders "
     "instantly and is the rail for refunds and bonuses, and every credit and debit is recorded as a "
     "transaction tied to its order or referral. (Top-up uses online payment, so it's done on the "
     "phone app, not the web.)",
     "A prepaid wallet makes checkout one-tap fast and gives the business a clean, instant way to "
     "issue refunds, credits, and referral bonuses.",
     "The wallet is only ever topped up after a verified payment server-side, and every movement is "
     "a logged, referenced transaction — so the balance is always reconcilable.")

flow(doc, "5.  Subscriptions",
     "A customer buys a plan (food or essentials) with a start date and pays in full upfront. After "
     "that, a scheduled job on the server creates that day's delivery order automatically from the "
     "plan's items — every day, at the right cycle — with no action from the customer. The plan "
     "winds down by meals actually delivered, not by the calendar, so a pause, a skipped day, or "
     "even a server outage simply extends the end date and the customer still receives every meal "
     "they paid for. Customers can pause or skip specific days; the app reminds them 1–2 days "
     "before the plan ends or if their wallet is low. An admin can cancel a subscription with an "
     "automatic prorated refund.",
     "It creates dependable recurring revenue while asking nothing daily of the customer — and it "
     "guarantees they never lose a paid meal to a pause or a glitch.",
     "“Days consumed, not dates” is the key idea: paid meals are always honored, and daily "
     "orders are generated server-side, so a subscription needs zero app interaction to keep "
     "flowing.")

flow(doc, "6.  Delivery — How Your Order Reaches You",
     "Every address is served one of two ways: direct (a delivery zone) or via a hub (a local "
     "handover point). In the hub model, a driver carries the order Dispatched → Received at Hub "
     "and hands off, then a hub operator takes it Received at Hub → On the Way → Delivered. "
     "In the direct model, the driver does the whole run. Drivers use a Driver Dashboard inside the "
     "customer app; hub operators use a Hub Dashboard. Each role can only make the moves that belong "
     "to it, and the customer is notified at each milestone.",
     "It supports both an own-fleet last mile and a partner-hub handover, with the correct person "
     "responsible for each leg.",
     "A single status engine splits responsibility by role, so no one can skip or reverse a step "
     "they don't own — keeping the delivery trail clean and accurate.",
     diagram=("Hub delivery hand-off:",
              ["Dispatched", "Received at Hub", "On the Way", "Delivered"], "38BDF8"))

flow(doc, "7.  Addresses, Serviceability, Zones & Hubs",
     "When a customer adds an address, they drop a map pin (Google Maps), and the server checks "
     "whether that point falls inside a delivery zone or hub area and, if so, marks the address "
     "serviceable and tags it with the right zone, hub, and branch. Someone outside the area can "
     "still “enter anyway” to browse, but they're nudged and blocked from checkout until "
     "they add a serviceable address. The delivery fee is resolved in priority order: hub override, "
     "then zone override, then the store default.",
     "The app only promises delivery where it can actually serve, and each order is automatically "
     "routed to the correct hub or zone.",
     "All the map math is done server-side, so the same rule governs the app, the daily dispatch, "
     "and the reports — no chance of the phone and the back-office disagreeing about who can be "
     "served.")

flow(doc, "8.  Referrals, Loyalty, Feedback & Ratings",
     "Every customer has a personal referral code and a shareable link. When a new customer applies "
     "a code, the server credits the new customer a sign-up wallet bonus (and optional points), and "
     "rewards the referrer when their referee places a first order or stays subscribed — and only "
     "once, so codes can't be farmed. Reward amounts are admin-set, not hard-coded. Loyalty points "
     "accrue on the profile (earn rate is admin-set) and can be redeemed. For quality, customers can "
     "leave overall feedback (rating + comment, optionally tied to an order) and rate individual "
     "items in an order; admins review all of it in one place.",
     "Referrals grow the customer base at almost no acquisition cost, loyalty points encourage "
     "repeat orders, and per-item ratings tell the kitchen precisely which dish needs attention.",
     "Every reward number is tunable by the admin at runtime, the credit logic is idempotent (no "
     "farming), and ratings are captured per dish, not just per order — so quality signals are "
     "specific and actionable.")

# ════════════ PART 2 — STAFF & OPERATIONS ════════════
part(doc, "Part 2 — For Staff & Operations")

flow(doc, "9.  Staff Order Processing — Kitchen & Packing (the Batch Board)",
     "Staff never chase individual orders. At each cycle's cutoff, the system “pushes” "
     "that cycle's batch to the staff screens — and the dashboard shows exactly one cycle's batch at "
     "a time, flipping to the next cycle when that one pushes, so two runs are never mixed. The "
     "Kitchen tab shows a combined prep list: the server totals every order's items so the kitchen "
     "cooks by quantity (e.g. “Item A × 12”), and staff advance the batch "
     "Confirmed → Ready. The Packing tab lists orders for labeling and advances "
     "Ready → Packed → Dispatched, with printable labels and summaries. Any past-dated "
     "order not yet delivered stays visible, so a missed delivery is never hidden by the batch flip.",
     "A single, current batch keeps the kitchen and packers focused on one run with a clear, "
     "aggregated to-cook list — far less error-prone than a long, mixed list of every open order.",
     "The board is driven by the cycle push itself, so it self-advances on the operational clock; "
     "the kitchen sees totals, not order-by-order, and undelivered orders are deliberately carried "
     "over so nothing perishable slips through.")

flow(doc, "10.  Staff Attendance (with Leave & Corrections)",
     "Staff clock in and out from their dashboard; each punch records the GPS location (via the "
     "phone's location service) and creates one record per person per day — a punch-in with no "
     "punch-out simply reads as “Present” for that day. Staff can view their month's "
     "history, request leave (an admin approves), and file an attendance correction to back-fill "
     "missed days with a reason, which an admin approves or rejects. All of this works offline: with "
     "no signal, the clock action is saved on the device and synced automatically when the "
     "connection returns.",
     "Attendance feeds payroll, so it must be reliable even in a kitchen with poor signal, and GPS "
     "on each punch discourages clocking in for someone else. Corrections give an honest, reviewable "
     "way to fix mistakes instead of silent edits.",
     "The offline queue is conflict- and identity-aware — on a shared device it won't replay one "
     "person's punch under another's account, and it won't double-apply. Most small-business apps "
     "simply lose attendance when the network drops.")

flow(doc, "11.  Internal Supply Orders (Veg / Grocery / Stationery) & Expense Claims",
     "Staff raise supply requests from their dashboard — vegetables, grocery, or stationery. Each "
     "request auto-merges into the admin's single active purchase list (the same item raised by two "
     "staff just adds up the quantity) while leaving an audit trail; there is no separate approval "
     "step — the admin editing the list is the approval. In Stock Manager the admin tweaks "
     "quantities, adds items (with name autocomplete), and prints a batch — which snapshots the "
     "current list as a PDF and clears the active list for the next cycle; past batches can be "
     "reprinted. Separately, staff file expense claims (amount, category, note); an admin approves "
     "and marks them paid.",
     "It turns scattered “we're out of onions” messages into one consolidated, printable "
     "purchase order, and gives expense reimbursement a tracked, auditable path.",
     "Requests deduplicate and total themselves automatically into a clean print-ready sheet — no "
     "spreadsheet, no approval bottleneck — yet every original request is preserved for audit.")

flow(doc, "12.  Offline Mode for Staff",
     "Staff actions that can't wait — order status updates and attendance punches — are queued on "
     "the device when offline and synced in order when the network returns. Two safeguards apply on "
     "sync: an update won't roll back an order someone else has already advanced, and on a shared "
     "device it won't replay one staffer's action under another's account.",
     "Kitchens and delivery routes have unreliable signal, so work must never be silently lost or "
     "mis-applied.",
     "The replay is conflict-aware and identity-safe — well beyond a simple “try again "
     "later” — protecting data integrity on the shared, low-signal devices these teams use.")

# ════════════ PART 3 — ADMIN & PLATFORM ════════════
part(doc, "Part 3 — Admin & the Platform")

flow(doc, "13.  Admin Controls — Every “Manage” Line, and Why",
     "The admin home has two tabs. Reports shows live drill-downs — Orders, Revenue, Subscriptions, "
     "Staff, Hub Delivery — all calculated on the server and visible only to admins. Manage is a "
     "searchable list of tools; each row and its purpose:",
     "Almost every business rule — prices, tax, cancellation windows, reward sizes, even the words "
     "in a push message — lives in the database, so the business can change it without a developer "
     "or an app update.",
     "The split between a normal branch admin and an all-seeing super-admin is enforced by the "
     "database itself, not just hidden in the screen — so the controls are genuinely safe, not "
     "cosmetic.",
     sub_bullets=[
        ("Manage Running Orders — ", "act on live orders (advance status, open details)."),
        ("Manage Running Subscriptions — ", "view active subs and cancel one with an automatic prorated refund."),
        ("Menu Manager — ", "add/edit food items and assign them to a delivery cycle."),
        ("Essentials Manager — ", "manage the grocery catalogue."),
        ("Subscriptions Manager — ", "define plans (items, duration, price)."),
        ("Delivery Manager — ", "draw and edit delivery zones and hubs on a map, assign drivers, set fees."),
        ("Note to Staff — ", "post a message banner onto the staff dashboard."),
        ("Manage Notifications — ", "edit the wording of every push message and switch each on or off."),
        ("Banners & Backgrounds — ", "set the home banner, login background, and hero image."),
        ("Referral Settings — ", "set referral reward amounts and toggle the program."),
        ("Customer Feedback — ", "read ratings and comments."),
        ("Resource Manager — ", "the people hub: staff list, onboard, salary/leave/attendance, offboard."),
        ("Expense Manager — ", "review/approve staff expense claims and record business expenses."),
        ("Stock Manager — ", "the internal purchase list (see #11) with printable batches."),
        ("Operations Manager — ", "master config: fees, tax, cancellation window, wallet limits, loyalty rate, storm mode, module switches, job health."),
        ("Manage Branches / Export Customers — ", "super-admin only."),
     ])

flow(doc, "14.  Automated Jobs — the Operational Clock",
     "A scheduled job runs every minute on the server: for each delivery cycle whose cutoff has just "
     "passed, it generates that day's subscription orders and pushes a kitchen summary to staff — "
     "and that push is also what “releases” the day's batch to the staff screens (see #9). "
     "Other jobs run daily and weekly: subscription-expiry reminders, low-wallet warnings, and a "
     "win-back nudge to customers who've gone quiet. All run on the server and need no one to press "
     "a button.",
     "The operational rhythm runs itself — the kitchen gets its list on time, customers get timely "
     "reminders, and lapsed users get re-engaged automatically.",
     "The whole cadence is server-driven and idempotent (safe to re-run), so nothing depends on a "
     "phone being open or a person remembering to trigger it.")

flow(doc, "15.  Notifications",
     "Push notifications are sent through Expo's push service. Every message resolves an "
     "admin-editable template by event (e.g. “order confirmed”, “order ready”), "
     "fills in the details, and is skipped entirely if the admin has switched that message off. "
     "Order milestones notify the customer; “Packed” is intentionally silent to avoid "
     "spamming. Every send is logged, and dead device tokens are automatically retired.",
     "Customers stay informed at the moments that matter, and admins control the wording without "
     "touching code.",
     "One template system governs every push across the app and the scheduled jobs, with per-message "
     "on/off switches and a full audit log of what was sent.")

flow(doc, "16.  Operations Config & Feature Flags",
     "A single operations record holds the store's economic settings — delivery fee, tax rate, "
     "cancellation window, wallet limits, loyalty earn rate, and warning thresholds. A set of "
     "feature flags turns whole modules on or off: the essentials (grocery) module, the referral "
     "program, hub delivery, multi-branch (fully built but currently switched off — the app runs as "
     "a single branch today), and storm mode, which instantly blocks new orders and renewals.",
     "The business can retune its economics, switch features on or off, and pause the entire store "
     "in bad weather — all in real time, with no app release.",
     "Launch-gated multi-branch and kill-switches mean the very same build scales from one branch to "
     "many and can be throttled instantly when conditions demand it.")

flow(doc, "17.  Where the App Runs — Phone, Web App & Website",
     "1stOne is one product that runs in three places. On Android and iOS phones it's the full "
     "experience (built with React Native / Expo, with over-the-air updates so fixes reach phones "
     "without a store release). The very same app also runs in a web browser as a Web App (via "
     "React Native for Web) — useful for customers without the app installed and especially for "
     "admins, who use the browser for heavy tasks like bulk CSV import and export of catalogue and "
     "customers. Because the online-payment SDK (Razorpay) is mobile-only, the web app accepts "
     "wallet payments only — adding money or paying online is done on the phone. Separately, there "
     "is a public marketing website at 1stone.in (static pages: home, FAQ, and a self-serve "
     "account-deletion page), built for search engines and link previews, that introduces the "
     "service and points people to the app.",
     "One shared codebase across phone and web means features and fixes land everywhere at once; the "
     "website gives the business a public, search-friendly front door and satisfies app-store "
     "requirements such as a self-serve account-deletion page.",
     "Customers and admins can pick the device that suits the task — phone for ordering and "
     "payments, browser for heavy admin work like CSV imports — without a second, separate product "
     "to maintain.")

doc.save(OUT)
print("wrote", OUT)
