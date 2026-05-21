#!/usr/bin/env python3
"""
1stOne F1 — branded .docx generator for "Document 2: Maintenance & Operations Runbook".

Same dark-theme branding as Document 1 (cover wordmark, TOC, page breaks),
plus checklist and data-table helpers for a runbook. Layered: Part 1 is an
owner checklist in plain English; Parts 2-8 are the technical reference.

NO secret values are printed — credentials are referenced by NAME and storage
location only. Content is factual and code-derived (mirrors 02-maintenance-runbook.md).
Run:  python3 docs/build_maintenance_docx.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
LOGO = os.path.join(HERE, "_logo_real.png")
if not os.path.exists(LOGO):
    LOGO = os.path.join(ROOT, "assets", "icon.png")
OUT = os.path.join(HERE, "02-maintenance-runbook.docx")

SKY        = RGBColor(0x38, 0xBD, 0xF8)
MINT       = RGBColor(0x4E, 0xCD, 0xC4)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
FONT       = "Tahoma"
BODY       = RGBColor(0xEC, 0xEC, 0xEC)
LIGHTGREEN = RGBColor(0xC5, 0xE0, 0xA0)
DARKTEXT   = RGBColor(0x22, 0x22, 0x22)
FOOTGREY   = RGBColor(0xAA, 0xAA, 0xAA)
AMBER      = RGBColor(0xFF, 0xBF, 0x00)
CONTENT_W  = 6.6

# ── low-level helpers ────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_bottom_border(cell, hexc="3A3A3A"):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement("w:tcBorders"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "4"); bot.set(qn("w:space"), "0"); bot.set(qn("w:color"), hexc)
    b.append(bot); tcPr.append(b)

def set_page_bg(doc, hex_color):
    bg = OxmlElement("w:background"); bg.set(qn("w:color"), hex_color)
    doc.element.insert(0, bg); doc.settings.element.append(OxmlElement("w:displayBackgroundShape"))

def set_update_fields(doc):
    el = OxmlElement("w:updateFields"); el.set(qn("w:val"), "true"); doc.settings.element.append(el)

def set_outline_level(p, level):
    pPr = p._p.get_or_add_pPr(); ol = OxmlElement("w:outlineLvl"); ol.set(qn("w:val"), str(level)); pPr.append(ol)

def add_page_number_footer(section):
    p = section.footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("1stOne — Maintenance & Operations Runbook   ·   ")
    run.font.name = FONT; run.font.size = Pt(8); run.font.color.rgb = FOOTGREY
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    r = p.add_run(); r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = FOOTGREY
    r._r.append(b); r._r.append(instr); r._r.append(e)

def style_base(doc):
    n = doc.styles["Normal"]; n.font.name = FONT; n.font.size = Pt(10.5); n.font.color.rgb = BODY
    n.paragraph_format.space_after = Pt(6); n.paragraph_format.line_spacing = 1.12

# ── building blocks ──────────────────────────────────────────────
def part(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(18); run.font.color.rgb = MINT
    pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "14"); bot.set(qn("w:space"), "2"); bot.set(qn("w:color"), "4ECDC4")
    pbdr.append(bot); pPr.append(pbdr); set_outline_level(p, 0); return p

def heading(doc, text, size=14, color=SKY, space_before=14, outline=1):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text); run.font.name = FONT; run.font.size = Pt(size); run.font.color.rgb = color
    if outline is not None: set_outline_level(p, outline)
    return p

def subhead(doc, text, color=LIGHTGREEN):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text); r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = color; return p

def rule(doc):
    p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6"); bot.set(qn("w:space"), "1"); bot.set(qn("w:color"), "38BDF8")
    pbdr.append(bot); pPr.append(pbdr); p.paragraph_format.space_after = Pt(2)

def body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs: p.runs[0].font.name = FONT; p.runs[0].font.color.rgb = BODY
    return p

def bullets(doc, items):
    for item in items:
        if isinstance(item, tuple): label, rest = item
        else: label, rest = "", item
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(2)
        if label:
            r = p.add_run(label); r.font.name = FONT; r.bold = True; r.font.color.rgb = LIGHTGREEN
        r2 = p.add_run(rest); r2.font.name = FONT; r2.font.color.rgb = BODY

def checklist(doc, items, mark="☐  "):
    for it in items:
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.18); p.paragraph_format.first_line_indent = Inches(-0.18)
        m = p.add_run(mark); m.font.name = FONT; m.font.size = Pt(11); m.font.color.rgb = MINT
        r = p.add_run(it); r.font.name = FONT; r.font.size = Pt(10.5); r.font.color.rgb = BODY

def note(doc, label, text, kind="rule"):
    fill = "2A2620" if kind == "warn" else "10242B"
    bar  = "FFBF00" if kind == "warn" else "38BDF8"
    lblcol = AMBER if kind == "warn" else SKY
    tbl = doc.add_table(rows=1, cols=1); tbl.autofit = True
    cell = tbl.cell(0, 0); set_cell_bg(cell, fill)
    tcPr = cell._tc.get_or_add_tcPr(); borders = OxmlElement("w:tcBorders"); left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "18"); left.set(qn("w:space"), "0"); left.set(qn("w:color"), bar)
    borders.append(left); tcPr.append(borders)
    cell.paragraphs[0].text = ""
    lr = cell.paragraphs[0].add_run(label.upper()); lr.font.name = FONT; lr.font.size = Pt(8); lr.bold = True; lr.font.color.rgb = lblcol
    bp = cell.add_paragraph(text); bp.runs[0].font.name = FONT; bp.runs[0].font.size = Pt(9.5); bp.runs[0].font.color.rgb = BODY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def data_table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False; tbl.allow_autofit = False
    for i, w in enumerate(widths):
        tbl.columns[i].width = Inches(w)
    # header
    for i, h in enumerate(headers):
        c = tbl.cell(0, i); c.width = Inches(widths[i]); set_cell_bg(c, "1F6F94"); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
        r = p.add_run(h); r.font.name = FONT; r.font.size = Pt(8.5); r.bold = True; r.font.color.rgb = WHITE
    # body
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            c = tbl.cell(ri, ci); c.width = Inches(widths[ci]); c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_bottom_border(c)
            p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
            r = p.add_run(val); r.font.name = FONT; r.font.size = Pt(8); r.font.color.rgb = BODY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_toc(doc):
    p = doc.add_paragraph(); run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    run._r.append(b); run._r.append(instr); run._r.append(sep)
    ph = p.add_run("Right-click here and choose “Update Field” to build the contents.")
    ph.font.name = FONT; ph.font.size = Pt(9.5); ph.italic = True; ph.font.color.rgb = FOOTGREY
    er = p.add_run(); e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end"); er._r.append(e)

# ════════════════════════════════════════════════════════════════
doc = Document(); style_base(doc)

# ---- COVER ----
sec = doc.sections[0]
sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9); sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)
set_page_bg(doc, "151515"); set_update_fields(doc)
for _ in range(4): doc.add_paragraph()
if os.path.exists(LOGO):
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER; lp.add_run().add_picture(LOGO, width=Inches(4.2))
doc.add_paragraph()
t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2r = t2.add_run("Maintenance & Operations Runbook"); t2r.font.name = FONT; t2r.font.size = Pt(26); t2r.font.color.rgb = WHITE
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("How to keep 1stOne healthy — what to check, when, and what never to break")
sr.font.name = FONT; sr.font.size = Pt(12); sr.font.color.rgb = MINT
for _ in range(6): doc.add_paragraph()
note_p = doc.add_paragraph(); note_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note_p.add_run("Document 2 of a series   ·   Reflects app v1.3.2-stable.1   ·   May 2026")
nr.font.name = FONT; nr.font.size = Pt(9); nr.font.color.rgb = FOOTGREY

# ---- BODY ----
doc.add_section(WD_SECTION.NEW_PAGE)
bsec = doc.sections[1]
bsec.top_margin = Inches(0.8); bsec.bottom_margin = Inches(0.8); bsec.left_margin = Inches(0.9); bsec.right_margin = Inches(0.9)
add_page_number_footer(bsec)

prov = doc.add_paragraph()
pr = prov.add_run("Prepared from a direct reading of the application's source code and configuration. "
                  "Where the code and older comments disagreed, the code was treated as the source of truth. "
                  "No secret values appear in this document — credentials are named by their variable and storage location only.")
pr.font.name = FONT; pr.font.size = Pt(9.5); pr.italic = True; pr.font.color.rgb = LIGHTGREEN
prov.paragraph_format.space_after = Pt(12)

ch = doc.add_paragraph(); chr_ = ch.add_run("Contents"); chr_.font.name = FONT; chr_.font.size = Pt(18); chr_.font.color.rgb = MINT
ch.paragraph_format.space_after = Pt(4); add_toc(doc)

doc.add_page_break()
heading(doc, "How to use this runbook", size=15, color=SKY, space_before=2, outline=1)
body(doc,
     "This runbook is layered. Part 1 is a plain-English checklist for the business owner / operator — "
     "no technical skill needed; tick the items on schedule and hand anything marked “(Tech)” "
     "to your developer. Parts 2–8 are the technical reference: monitoring, the dependency and "
     "credential register, deployment rules, security, backups, dos & don'ts, and code-derived "
     "watch-items. Everything is drawn from the actual code and configuration.")
note(doc, "Golden daily habit",
     "Open Admin → Operations Manager → Job Health every operating day. If every background job shows "
     "“succeeded” with zero 24-hour failures, the engine that cooks, dispatches, charges, and "
     "notifies is healthy. This single screen is your most important check.")

# ════════════ PART 1 — OWNER CHECKLIST ════════════
part(doc, "Part 1 — Owner's Maintenance Checklist")
body(doc, "No technical skill needed for this part. Work top-to-bottom on each schedule; "
          "delegate items marked “(Tech)”.")

subhead(doc, "Every operating day")
checklist(doc, [
    "Open Admin → Operations Manager → Job Health: every background job shows “succeeded”, no 24-hour failures.",
    "Glance at the Reports tab — today's Orders and Revenue look sane.",
    "Act on any reconciliation alert push: the app notifies admins if a wallet refund failed or a subscription was paid but not created.",
    "Confirm the kitchen received its summary push at each cycle's cutoff (ask staff, or check Job Health's recent manifest runs).",
])
subhead(doc, "Every week")
checklist(doc, [
    "Job Health → push outcomes: confirm expiry reminders, low-wallet warnings, and the Monday win-back went out.",
    "Review Customer Feedback and item ratings; route recurring complaints to the kitchen.",
    "(Tech) Skim Sentry for new app errors/crashes.",
    "(Tech) Confirm a recent database backup exists in Supabase.",
])
subhead(doc, "Every month")
checklist(doc, [
    "Run staff salary; reconcile attendance, leave, and expense claims (Resource Manager / Expense Manager).",
    "Review store settings — delivery fee, tax, cancellation window, wallet limits, loyalty rate — still correct?",
    "Review feature flags — anything to switch on or off?",
    "(Tech) Apply dependency and security updates; re-run the test gate.",
])
subhead(doc, "Every quarter")
checklist(doc, [
    "(Tech) Rotate sensitive secrets: Razorpay key-secret & webhook secret, Supabase service-role key.",
    "Review who holds admin / super-admin access; offboard anyone who has left.",
    "(Tech) Plan any Expo SDK or major library upgrades.",
])
subhead(doc, "Every year (set calendar reminders — see Part 3 register)")
checklist(doc, [
    "Renew the Apple Developer Program membership — if it lapses, the iOS app is removed from sale.",
    "Renew the domain 1stone.in with the registrar.",
    "Re-check Google Play and Apple policy / target-SDK requirements and update if required.",
    "Confirm the payment gateway (Razorpay) account / KYC is active and any plan is renewed.",
])

# ════════════ PART 2 — HEALTH & MONITORING ════════════
part(doc, "Part 2 — Health Checks & Monitoring")
heading(doc, "Where to look", size=14); rule(doc)
bullets(doc, [
    ("In-app Job Health — ", "Admin → Operations Manager → Job Health (the get_job_health RPC, admin-only). Shows every scheduled job's last run, status, and 24-hour failure count; recent subscription-dispatch (manifest) runs; and push-delivery outcomes over 24h."),
    ("Cron-failure alert — ", "an hourly job pushes admins automatically if any scheduled job failed — your early-warning even before you open Job Health."),
    ("Supabase dashboard — ", "Edge Function logs, Postgres logs, and cron.job_run_details for deep dives."),
    ("Reconciliation alerts — ", "the order/cancel functions push branch admins on a wallet-refund failure or a paid-but-not-created subscription. These need a human to fix."),
    ("Sentry — ", "app crashes and errors. PostHog — product analytics. Razorpay dashboard — payments, settlements, webhook delivery."),
])
subhead(doc, "Tables that record what happened")
bullets(doc, [
    ("manifest_run_log — ", "each daily subscription-dispatch run (orders created / skipped, errors)."),
    ("kitchen_push_log — ", "each cycle's kitchen push (also drives the staff “active batch”)."),
    ("push_logs — ", "every notification attempt, with delivery status."),
])
subhead(doc, "Scheduled jobs (the operational clock)")
data_table(doc,
    ["Job", "Schedule (UTC)", "What it does"],
    [
        ["kitchen-cutoff-push-tick", "* * * * * (every min)", "At each cycle's cutoff: generate that day's subscription orders + push the kitchen summary (sets the staff batch)."],
        ["subscription-expiry-push", "30 3 * * *  (09:00 IST)", "1-day / 2-day expiry reminders + “starts tomorrow” notices."],
        ["low-wallet-check", "0 4 * * *  (09:30 IST)", "Warn customers with a low wallet before a subscription auto-renews."],
        ["dormant-user-check", "30 4 * * 1  (Mon 10:00 IST)", "Weekly win-back nudge to customers who've gone quiet."],
        ["expire-idempotency-keys", "0 * * * *  (hourly)", "Clean up old idempotency rows."],
        ["cron-failure-alert", "15 * * * *  (hourly)", "Push admins if any scheduled job failed."],
    ],
    [1.9, 1.55, 3.15])

# ════════════ PART 3 — DEPENDENCIES, ACCOUNTS, CREDENTIALS ════════════
part(doc, "Part 3 — Dependencies, Accounts, Credentials & Renewals")
heading(doc, "Dependency inventory", size=14); rule(doc)
data_table(doc,
    ["Layer", "Technology", "Version", "Notes"],
    [
        ["App framework", "Expo SDK", "~54", "Build tooling + over-the-air (OTA) updates"],
        ["Mobile runtime", "React Native", "0.81.5", "Hermes engine"],
        ["UI runtime", "React / React DOM", "19.1.0", "Also powers the web app"],
        ["Navigation", "@react-navigation", "7", "Native stack per role"],
        ["Server state", "TanStack Query", "5", "Caching / fetching"],
        ["Local state", "Zustand", "5", "Cart, UI, offline queue (AsyncStorage)"],
        ["Backend client", "@supabase/supabase-js", "2.45", "App + edge functions"],
        ["Payments", "react-native-razorpay", "2.3 (patched)", "patches/ patch must persist (postinstall)"],
        ["Maps", "react-native-maps / @react-google-maps/api", "1.20 / 2.20", "Needs Google Maps key"],
        ["Notifications", "expo-notifications", "—", "Push via Expo"],
        ["Monitoring", "@sentry/react-native + posthog", "8.x / 4.x", "Errors + analytics"],
        ["Edge runtime", "Deno (Supabase Edge)", "std 0.168", "17 functions; supabase-js@2.45"],
        ["DB extensions", "pg_cron, pg_net, Vault", "—", "Scheduling, HTTP calls, secret storage"],
        ["Build / CI", "EAS", "cli >= 12", "Profiles: development / preview / production"],
        ["Dev tooling", "TypeScript / ESLint / Jest / Husky / Knip", "5.9 / 9 / 30 / 9", "Pre-push gate: tsc + jest"],
    ],
    [1.35, 2.35, 1.1, 1.8])

heading(doc, "External accounts, credentials & renewals", size=14); rule(doc)
body(doc, "Secrets live in .env (local, git-ignored), EAS build env / secrets, and the Supabase "
          "Edge Function environment + Vault. This document names them only. Fill the last column "
          "with the account owner and next renewal/rotation date.")
data_table(doc,
    ["Service", "Credential name & where it lives", "Rotation / Renewal", "Owner & next date (fill in)"],
    [
        ["Supabase (project wcvqxzqqwcxlcgrjyunf)", "EXPO_PUBLIC_SUPABASE_ANON_KEY (public); SUPABASE_SERVICE_ROLE_KEY (server only — function env + Vault + app_config)", "Rotate service key on staff change / leak", "____________"],
        ["Razorpay", "EXPO_PUBLIC_RAZORPAY_KEY_ID; RAZORPAY_KEY_SECRET; RAZORPAY_WEBHOOK_SECRET (secrets in function env)", "Rotate secret + webhook periodically; keep KYC active", "____________"],
        ["Google Maps", "EXPO_PUBLIC_GOOGLE_MAPS_KEY (Android manifest)", "Usage-billed; keep key restricted", "____________"],
        ["Firebase / FCM", "google-services.json (Android)", "Re-download if project changes", "____________"],
        ["Expo / EAS", "Project 81ff7f3c-… ; Expo account; OTA via u.expo.dev", "Account active; submission creds", "____________"],
        ["Apple App Store", "Bundle com.1stone.f1; appleId / ascAppId / teamId (eas.json — currently placeholders)", "Developer Program ≈ yearly (~$99)", "____________"],
        ["Google Play", "Package com.stone1st.f1; play-store-service-account.json", "One-time reg; policy/target-API yearly", "____________"],
        ["Domain", "1stone.in", "Registrar — yearly", "____________"],
        ["Sentry", "EXPO_PUBLIC_SENTRY_DSN (publishable)", "Plan tier", "____________"],
        ["PostHog", "EXPO_PUBLIC_POSTHOG_KEY / _HOST", "Plan tier", "____________"],
    ],
    [1.45, 2.55, 1.45, 1.15])

# ════════════ PART 4 — DEPLOYMENT & RELEASE ════════════
part(doc, "Part 4 — Deployment & Release")
note(doc, "Owner note",
     "Most bug fixes reach phones automatically as over-the-air (OTA) updates — no app-store wait. "
     "Only changes to the phone's native parts (new permissions, SDK upgrades, new native libraries) "
     "need a full rebuild and store review, which is slower.")
heading(doc, "Two release paths", size=14); rule(doc)
bullets(doc, [
    ("OTA update (JS / content only) — ", "ship JS logic, copy, and bug fixes with an Expo update; phones pick it up on next launch (checkAutomatically = ON_LOAD). No store review. Valid only within the same Expo SDK (runtimeVersion = sdkVersion)."),
    ("Native build — ", "for new native code, an SDK upgrade, new permissions, or a version bump: build with EAS (production profile; autoIncrement on) and submit to the stores."),
])
heading(doc, "Backend deploys", size=14); rule(doc)
bullets(doc, [
    ("Edge Functions — ", "supabase functions deploy <name> (most use --no-verify-jwt). Critical: deploy place-order together with the matching app build — its request contract is not backward-compatible."),
    ("SQL — ", "files in supabase/sql are idempotent; run them in the order in DEPLOY_SQL_ORDER.md. Add new migration files; never hand-edit old ones. After the access-token hook is installed, enable it in the dashboard."),
    ("Prerequisites — ", "extensions pg_cron + pg_net; Vault secrets supabase_url + service_role_key; the app_config rows used for in-database push must exist."),
])
note(doc, "Quality gate",
     "A Husky pre-push hook runs tsc --noEmit and jest; the push is blocked if either fails. "
     "Run it yourself first with: npm run check.")

# ════════════ PART 5 — SECURITY & SECRET ROTATION ════════════
part(doc, "Part 5 — Security & Secret Rotation")
heading(doc, "Which secrets exist", size=14); rule(doc)
bullets(doc, [
    ("Sensitive (never expose) — ", "Supabase service-role key, Razorpay key-secret, Razorpay webhook secret. These live server-side only (Supabase function env / Vault), never in the app bundle."),
    ("Publishable (safe in the app) — ", "Supabase anon key, Google Maps key, Sentry DSN, PostHog key, Razorpay key-id. These are embedded in builds by design."),
])
heading(doc, "Rotation steps", size=14); rule(doc)
bullets(doc, [
    ("Supabase service-role key — ", "roll it in the dashboard, then update it everywhere it is stored: the Edge Function env, the Vault secret, and the app_config row used for in-database push. Redeploy functions."),
    ("Razorpay secret / webhook — ", "rotate in the Razorpay dashboard, update the function env and the webhook configuration, then redeploy."),
])
heading(doc, "Standing rules", size=14); rule(doc)
bullets(doc, [
    ("Row-Level Security stays on — ", "every table is RLS-protected; money and role changes flow only through SECURITY DEFINER atomic RPCs or service-role functions. Never write those columns directly."),
    ("Never disable the webhook signature check — ", "verify-payment verifies an HMAC signature; without it, anyone could mark orders paid."),
    ("Access review — ", "admin / super-admin is set on the profile; remove leavers and offboard staff through the proper RPC."),
])

# ════════════ PART 6 — BACKUPS & DR ════════════
part(doc, "Part 6 — Backups & Disaster Recovery")
heading(doc, "What to protect", size=14); rule(doc)
bullets(doc, [
    ("Database — ", "rely on Supabase automated backups; confirm your plan's retention / point-in-time recovery, and periodically test a restore."),
    ("Android signing keystore — ", "managed by EAS. Losing the signing key means you can no longer ship updates to the same Play listing — ensure EAS holds it (or keep a secure backup)."),
    ("Apple distribution certificate / profiles — ", "managed via EAS; keep Apple account access safe."),
    ("Secrets (.env and server secrets) — ", "keep a secure copy (e.g. a password manager); they are not in git."),
    ("Code & schema — ", "git is the source of truth; supabase/sql + DEPLOY_SQL_ORDER.md reconstruct the database; a schema.sql snapshot exists."),
    ("Storage bucket — ", "the assets bucket (logo, banners, PDFs) — back it up."),
])
note(doc, "Recovery drill",
     "You should be able to: (1) restore the database from a Supabase backup, (2) redeploy the edge "
     "functions, (3) rebuild the app from git, and (4) re-provision the secrets. Rehearse it before you need it.")

# ════════════ PART 7 — DOS & DON'TS ════════════
part(doc, "Part 7 — Dos & Don'ts")
subhead(doc, "Do", color=LIGHTGREEN)
checklist(doc, [
    "Run npm run check (tsc + jest) before pushing — the pre-push gate enforces it.",
    "Keep the patches/ folder; patch-package re-applies the Razorpay patch on every install.",
    "Deploy place-order (and any changed edge functions) together with the matching app build.",
    "Keep business values in store_config / feature_flags and styling in the theme — the codebase mandates zero hardcoded values.",
    "Follow DEPLOY_SQL_ORDER.md; add new SQL files instead of editing old ones.",
    "Test owner-facing flows on a real device before a release.",
    "Keep RLS on and route money / role changes through the atomic RPCs.",
], mark="✔  ")
subhead(doc, "Don't", color=AMBER)
checklist(doc, [
    "Don't weaken the server-authority, quote-drift, or idempotency rules in the order path.",
    "Don't hardcode prices, fees, hex colours, or fonts — use config and the theme.",
    "Don't commit secrets or paste the service-role / Razorpay secrets anywhere shared.",
    "Don't disable the Razorpay webhook signature verification.",
    "Don't hand-edit money / role columns or bypass RLS in the database.",
    "Don't delete idempotency_keys or the audit tables (push_logs, manifest_run_log, kitchen_push_log).",
    "Don't ship a native change as an OTA update — it won't take effect; rebuild instead.",
], mark="✘  ")

# ════════════ PART 8 — WATCH-ITEMS ════════════
part(doc, "Part 8 — Watch-items & Risks (from the code)")
body(doc, "These are factual observations from the current code and configuration — worth a decision, not necessarily a problem.")
note(doc, "Razorpay test key in the production build profile", "warn")
bullets(doc, [
    ("", "eas.json's production profile lists EXPO_PUBLIC_RAZORPAY_KEY_ID as a test key (rzp_test_…). Confirm the LIVE key id is used for production builds (e.g. injected via EAS secrets) before real payments."),
])
note(doc, "iOS submission not configured", "warn")
bullets(doc, [
    ("", "eas.json submit → iOS still has REPLACE_WITH_* placeholders for Apple ID / team / app id. Set these before an App Store submission."),
])
heading(doc, "Other items to track", size=14); rule(doc)
bullets(doc, [
    ("Multi-branch is built but OFF — ", "the branch_management_active flag is false (single-branch today). Turning it on needs a branch_id data backfill and users to refresh their login token."),
    ("Two essentials gates — ", "store_config.essentials_module_active and an essentials feature flag both exist; consolidate to avoid confusion."),
    ("Generated DB types lag the live schema — ", "several RPCs are called with type casts (“types not regenerated”). Run npm run supabase:gen-types after any schema change."),
    ("Service-role key is stored in three places — ", "function env, Vault, and the app_config table (for in-database push). Rotating the key means updating all three."),
    ("Storm mode is a kill switch — ", "the storm_mode_active flag instantly blocks all new orders and renewals; know it exists for incidents (and remember to turn it back off)."),
])

doc.save(OUT)
print("wrote", OUT)
